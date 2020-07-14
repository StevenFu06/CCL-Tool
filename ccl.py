"""CCL Module for docx CCL interface

Date: 2020-07-03
Revision: A
Author: Steven Fu
Last Edit: Steven Fu
"""

from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt

from package import Parent, _re_doc_num, _re_pn
from compare import Rearrange, Bom, Tracker
from filehandler import *

import pandas as pd
from docx.api import Document
import shutil


class CCL:
    """Main CCL Class

    This class interfaces/ ties together all modules into one easily callalble class.
    Class is then call by the gui to perform all functions.

    Attributes:
        ccl_docx (str): ccl document path
        filtered (dataframe): dataframe of the filtered dataframe
        avl_bom (dataframe): dataframe of the avl bom
        avl_bom_path (str): Path to the AVl Bom
        avl_bom_updated (dataframe): dataframe of an updated avl bom
        avl_bom_updated_path (str): path to the updated avl bom
        path_illustration (str): path to the illustration folder
        path_ccl_data (str): path to the CCL documents folder
        path_checks (str): paths to check before downloading form enovia
        username (str): Enovia username
        password (str): Enovia password
    """
    def __init__(self):
        # Files
        self.ccl_docx = None  # Docx path
        self.filtered = None  # df
        self.avl_bom = None  # df
        self.avl_bom_updated = None  # df
        # Save paths
        self.avl_bom_path = None
        self.avl_bom_updated_path = None
        self.path_illustration = None
        self.path_ccl_data = None
        self.path_checks = []
        # Enovia
        self.username = None
        self.password = None
        # Parallel Computing
        self.processes = 1

########################################################################################################################
# Bom comparison
########################################################################################################################

    def set_bom_compare(self, avl_bom_old: str, avl_bom_new: str):
        """Sets the bom vairables, convert file paths to df

        Parameters:
            :param avl_bom_old: filepath to the old avl multilevel bom
            :param avl_bom_new: filepath to the new avl multilevel bom
        """
        # Set old
        self.avl_bom_path = avl_bom_old
        self.avl_bom_updated_path = avl_bom_new
        # Set new
        self.avl_bom = CCL.read_avl(avl_bom_old, 0)
        self.avl_bom_updated = CCL.read_avl(avl_bom_new, 0)

    @staticmethod
    def read_avl(path: str, skiprow: int):
        """Read CSV and determine headers row

        Parameters:
            :param path: path to the csv file
            :param skiprow: recursive call to determine which row is header
        """
        df = pd.read_csv(path, skiprows=skiprow)
        try:
            df['Name']
        except KeyError:
            skiprow += 1
            if skiprow < 10:
                df = CCL.read_avl(path, skiprow)
            else:
                raise TypeError('File is in wrong format')
        return df

    def avl_path_to_df(self):
        """Converts avl_bom_path to df if not given"""

        self.avl_bom = pd.read_csv(self.avl_bom_path)
        self.avl_bom_updated = pd.read_csv(self.avl_bom_updated_path)

    def bom_compare(self):
        """Performs a bom compare

        tracker: contains the information of removed and updated parts
        tracker_reversed: contains the information of added parts

        :return: tracker, tracker_reversed
        """
        if self.avl_bom is None or self.avl_bom_updated is None:
            raise ValueError('Missing required fields, ccl, avl_new or avl_old')
        #  Create bom object for forward compare
        tree, bom, tree_updated, bom_updated = self._get_bom_obj()
        tracker = Tracker()
        Rearrange(bom, bom_updated, tracker)
        # Create new object because previous object was modified during rearraange process
        tree, bom, tree_updated, bom_updated = self._get_bom_obj()
        tracker_reversed = Tracker()
        Rearrange(bom_updated, bom, tracker_reversed)
        return tracker, tracker_reversed

    def _get_bom_obj(self):
        """Used to create bom object"""

        tree = Parent(self.avl_bom).build_tree()
        bom = Bom(self.avl_bom, tree)

        tree_updated = Parent(self.avl_bom_updated).build_tree()
        bom_updated = Bom(self.avl_bom_updated, tree_updated)
        return tree, bom, tree_updated, bom_updated

    def save_compare(self, save_name: str):
        """Outputs the BOM comparison to a nice format

        Parameters:
            :param save_name: Save name/ path of the zip file

        :returns: outputs a zip file containing a added.csv, changed.csv, and removed.csv
        """
        # Perform bom comparison
        tracker, tracker_reversed = self.bom_compare()
        # Creates temporary directory to be converted into a zip file
        path = os.path.join(os.getcwd(), 'bom compare temp')
        if not os.path.exists(path):
            os.makedirs(path)
        # Format changed.csv
        df_updated = tracker.combine_found().reset_index()
        changed = {'old index': [], 'old pn': [], 'old description': [],
                   'new index': [], 'new pn': [], 'new description': []}
        for idx in df_updated.index:
            changed['old index'].append(df_updated.loc[idx, 'old_idx'])
            changed['old pn'].append(self.avl_bom.loc[df_updated.loc[idx, 'old_idx'], 'Name'])
            changed['old description'].append(self.avl_bom.loc[df_updated.loc[idx, 'old_idx'], 'Description'])

            changed['new index'].append(df_updated.loc[idx, 'new_idx'])
            changed['new pn'].append(self.avl_bom_updated.loc[df_updated.loc[idx, 'new_idx'], 'Name'])
            changed['new description'].append(self.avl_bom_updated.loc[df_updated.loc[idx, 'new_idx'], 'Description'])
        pd.DataFrame.from_dict(changed).to_csv(os.path.join(path, 'changed.csv'))
        # Format removed.csv
        removed = {'Part Number': [], 'Description': []}
        for idx in tracker.not_found_to_df()['idx']:
            removed['Part Number'].append(self.avl_bom.loc[idx, 'Name'])
            removed['Description'].append(self.avl_bom.loc[idx, 'Description'])
        pd.DataFrame.from_dict(removed).to_csv(os.path.join(path, 'removed.csv'))
        # Format added.csv
        added = {'Part Number': [], 'Description': []}
        for idx in tracker_reversed.not_found_to_df()['idx']:
            added['Part Number'].append(self.avl_bom_updated.loc[idx, 'Name'])
            added['Description'].append(self.avl_bom_updated.loc[idx, 'Description'])
        pd.DataFrame.from_dict(added).to_csv(os.path.join(path, 'added.csv'))
        # Zip and cleanup
        shutil.make_archive(save_name.replace('.zip', ''), 'zip', path)
        shutil.rmtree(path)

########################################################################################################################
# CCL Updating
########################################################################################################################

    def update_ccl(self, save_path: str, ccl_docx: str = None):
        """Will perform a BOM comparison as well as a CCL Update

        Parameters:
            :param save_path: save location of the updated CCL
            :param ccl_docx: same as the class ccl_docx, CCL docx location

        :return: A new updated CCL saved to the specified location
        """

        if ccl_docx is not None:
            self.ccl_docx = ccl_docx
        elif self.ccl_docx is None:
            raise ValueError('CCL is not given')

        tracker = self.bom_compare()[0]
        all_updates, removed = tracker.combine_found(), tracker.not_found_to_df()
        ccledit = CCLEditor(self.ccl_docx)
        self._updates_only(ccledit, all_updates)
        self._removed_only(ccledit, removed)
        ccledit.save(save_path)

    def _updates_only(self, ccledit, all_updates):
        """Deals with updating the CCL only

        All changes made to the CCL is done through the CCLEdit object

        Parameters:
            :param ccledit: ccledit object for easy editing of the docx ccl
            :param all_updates: a dataframe including added and changed part numbers (tracker.combine_found())

        Calls _update_pn, _update_desc_fn, _update_manufacturer, _update_model, and match_conditions
        to update the formating given all_updates.

        The above functions have params:
            :param row: row to be updated
            :param to_update: dataframe of to be updated parts
            :param ccledit: CCL edit object
        """
        for row in range(len(ccledit.table.rows)):
            pn = ccledit.get_text(row, 0)
            if pn in all_updates['old_pn'].to_list():
                to_update = all_updates.loc[pn == all_updates['old_pn']].values[0]
                # previous formatting
                bold = ccledit.isbold(row, 0)
                # Update fields
                self._update_pn(row, to_update, ccledit)
                self._update_desc_fn(row, to_update, ccledit)
                self._update_manufacturer(row, to_update, ccledit)
                self._update_model(row, to_update, ccledit)
                if bold:
                    ccledit.bold_row(row)
                self._match_conditions(row, ccledit, to_update)
                # Remove updated from df
                all_updates = all_updates[all_updates.old_idx != to_update[0]]

    def _update_pn(self, row, to_update, ccledit):
        """Updates column 1, part number"""

        align = ccledit.get_justification(row, 0)
        ccledit.set_text(row, 0, to_update[3])
        ccledit.set_justification(row, 0, align)

    def _update_desc_fn(self, row, to_update, ccledit):
        """Updates description and find number"""

        desc = self.avl_bom_updated.loc[to_update[2], 'Description']
        fn = self.avl_bom_updated.loc[to_update[2], 'F/N']
        ccledit.set_text(row, 1, f'{desc} (#{fn})')

    def _update_manufacturer(self, row, to_update, ccledit):
        """Updates manufacturer"""

        # Pop manufacturers one at a time because multiple can exist
        manufacturer = self.avl_bom_updated.loc[to_update[2], 'Manufacturer'].split('\n')[0]
        self.avl_bom_updated.loc[to_update[2], 'Manufacturer'] = \
            self.avl_bom_updated.loc[to_update[2], 'Manufacturer'].replace(manufacturer + '\n', '')
        # if empty replace with AB sciex but also warn the user
        if manufacturer.isspace():
            manufacturer = 'AB Sciex'
            print(f'{to_update[3]} couldnt find manufacturer')
        ccledit.set_text(row, 2, manufacturer)

    def _update_model(self, row, to_update, ccledit):
        """Updates the model field in the CCL"""

        # Pop Equivalent one at a time because multiple can exist
        model = self.avl_bom_updated.loc[to_update[2], 'Equivalent'].split('\n')[0]
        self.avl_bom_updated.loc[to_update[2], 'Equivalent'] = \
            self.avl_bom_updated.loc[to_update[2], 'Equivalent'].replace(model + '\n', '')
        # if empty replace with part number only but also warn the user
        if model.isspace():
            model = to_update[3]
            print(f'{to_update[3]} couldnt find Equivalent')
        ccledit.set_text(row, 3, model)

    def _match_conditions(self, row, ccledit, to_update):
        """Will highlight/ format the changed row"""

        if to_update[4] == 'full':
            ccledit.highlight_row(row, 'BRIGHT_GREEN')
        elif to_update[4] == 'partial':
            ccledit.highlight_row(row, 'RED')
        elif to_update[4] == 'fn_only':
            ccledit.highlight_row(row, 'YELLOW')

    def _removed_only(self, ccledit, removed):
        """Similar _update_only, this only deals with the removed items

        Will format and edit the CCL through the CCLEdit object

        Parameter:
            :param ccledit: CCLEdit object
            :param removed: removed dataframe form tracker
        """
        for row in range(len(ccledit.table.rows)):
            pn = ccledit.get_text(row, 0)
            if pn in removed['pn'].to_list():
                ccledit.strike_row(row)

########################################################################################################################
# Specification Documents Gathering
########################################################################################################################

    def collect_documents(self, headless: bool=True):
        """Collects the documents given the CCL

        :param headless: Run selenium headless mode, default is yes

        :return: structured folder based according to CSA submission package
        """
        # Error Checking/ missing info
        if self.ccl_docx is None:
            raise ValueError('CCL document is not given')
        if self.username is None or self.password is None:
            raise ValueError('Enovia username or password not given')
        if self.path_ccl_data is None:
            raise ValueError('CCL Documents save location not given')

        # Calls document collector
        collector = DocumentCollector(username=self.username,
                                      password=self.password,
                                      ccl=self.ccl_docx,
                                      save_dir=self.path_ccl_data,
                                      processes=self.processes,
                                      headless=headless)
        collector.collect_documents(check_paths=self.path_checks)

########################################################################################################################
# Illustration gathering and CCL updating
########################################################################################################################

    def collect_illustrations(self):
        """Collects the illustrations using the Illustration class found in filehandler"""

        # Error checking
        if self.ccl_docx is None:
            raise ValueError('CCL document is not given')
        if self.path_illustration is None:
            raise ValueError('Illustration save path not given')

        # Collect ills
        illustration = Illustration(ccl=self.ccl_docx,
                                    save_dir=self.path_illustration,
                                    processes=self.processes)
        illustration.get_illustrations(ccl_dir=self.path_ccl_data)

    def insert_illustration_data(self, save_path: str):
        """Insert illustration data into CCL

        Will insert the illustration data according to the illustration folder into the CCL.
        Will overwrite any existing illustration data in the CCL

        :param save_path: Save location of updated CCL
        """
        if self.ccl_docx is None:
            raise ValueError('CCL Document is not given')
        ccledit = CCLEditor(self.ccl_docx)
        for row in range(len(ccledit.table.rows)):
            pn = _re_pn(ccledit.get_text(row, 0))
            bold = ccledit.isbold(row, 4)

            new_technical = self.new_illustration_data(pn) + self.remove_illustration_data(ccledit.get_text(row, 4))
            ccledit.set_text(row, 4, new_technical)

            if bold:
                ccledit.set_bold(row, 4)
        ccledit.save(save_path)

    def new_illustration_data(self, pn: int):
        """Format the techincal data to insert illustration data into the column

        :param pn: part number
        :return: a string with properly formatted illustration data and number
        concated with existing text.
        """
        # Getting the illustration data information
        info = []
        for file in os.listdir(self.path_illustration):
            if file.endswith('.pdf'):
                ill_num = file.split(' ')[0]
                file_pn = file.split(' ')[1]
                dnum = _re_doc_num(file)
                sch_assy = 'Assy.' if 'Assy.' in file.split(' ') else 'Sch.'
                if str(pn) == str(file_pn):
                    info.append((ill_num, dnum[0], sch_assy))
        # Foramtting and concatenating with existing data for final techincal data column
        illustration_data = 'Refer to'
        if info:
            for ill_num, dnum, sch_assy in info:
                illustration_data = illustration_data + f' {ill_num} {sch_assy} {dnum};'
            return illustration_data
        return ''

    @staticmethod
    def remove_illustration_data(technical_string: str):
        """Removes any illustration data/ reference from technical column

        :param technical_string: text extracted from technical data column

        :return: a cleaned technical string stripped of illustration data
        """
        # Regex expression to accommodate as much variation and user input variation as possible
        results = re.findall(r'(?:\s*,|and)?\s*(?:Refer to)?\s*(?:Ill.|Ill)\s*(?:\d+.|\d+)\s*'
                             r'(?:Sch.|Assy.|Sch|Assy)\s*D\d+\s*(?:;|and)?',
                             technical_string, re.IGNORECASE)
        for result in results:
            technical_string = technical_string.replace(result, '')
        return technical_string

    def insert_illustration(self, ill_num: int, new_ill: str, save_path: str):
        """Inserts an illustration and updates the CCL with the new illustration

        Parameters:
            :param ill_num: illustration number
            :param new_ill: new illustration location
            :param save_path: save path of the illustration folder
        """
        illustration = Illustration(self.ccl_docx, self.path_illustration)
        illustration.shift_up_ill(ill_num)
        copyfile(new_ill, self.path_illustration)
        self.insert_illustration_data(save_path)

    def delete_illustration(self, ill_num: int, rm_ill: str, save_path: str):
        """"Deletes an illustration and updates the CCL with the new illustration

        Parameters:
            :param ill_num: illustration number
            :param rm_ill: illustration to be removed location
            :param save_path: save path of the illustration folder
        """
        illustration = Illustration(self.ccl_docx, self.path_illustration)
        illustration.shift_down_ill(ill_num)
        os.remove(rm_ill)
        self.insert_illustration_data(save_path)


class CCLEditor:
    """CCL Editor for easy interfacing with Python-docx and the CCL

    Attributes:
        document: the word document containing the ccl
        table: the table within the word document containing the CCL

    All functions require a row number for the row to be modified.
    Funcitons ending with XXXX_row means the entire row will be modified,
    while others will modify only modify one specific cell.
    """
    def __init__(self, docx_path):
        self.document = Document(docx_path)
        self.table = self.document.tables[0]

    def get_text(self, row, column):
        return self.table.rows[row].cells[column].text

    def set_text(self, row, column, new_string):
        self.table.rows[row].cells[column].text = new_string

    def set_bold(self, row, column):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    def bold_row(self, row):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    def isbold(self, row, column):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                if run.font.bold:
                    return True
        return False

    def set_italic(self, row, column):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.italic = True

    def italic_row(self, row):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.italic = True

    def isitalic(self, row, column):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                if run.font.italic:
                    return True
        return False

    def set_highlight(self, row, column, colour):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.highlight_color = getattr(WD_COLOR_INDEX, colour)

    def highlight_row(self, row, colour):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.highlight_color = getattr(WD_COLOR_INDEX, colour)

    def set_colour(self, row, column, r, g, b):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(r, g, b)

    def colour_row(self, row, r, g, b):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(r, g, b)

    def get_font(self, row, column):
        return self.table.rows[row].cells[column].paragraphs[0].runs[0].font.name

    def set_font(self, row, column, font):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.name = font

    def font_row(self, row, font):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = font

    def set_fontsize(self, row, column, size):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(size)

    def fontsize_row(self, row, size):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)

    def get_justification(self, row, column):
        return self.table.rows[row].cells[column].paragraphs[0].alignment

    def set_justification(self, row, column, justification):
        # just = {'left': 0, 'center': 1, 'right': 3, 'distribute': 4}
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            paragraph.alignment = justification

    def set_strike(self, row, column):
        for paragraph in self.table.rows[row].cells[column].paragraphs:
            for run in paragraph.runs:
                run.font.strike = True

    def strike_row(self, row):
        for cell in self.table.rows[row].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.strike = True

    def save(self, path):
        self.document.save(path)


if __name__ == '__main__':
    import pandas as pd
    ccl = CCLEditor('rev c bugatti.docx')
    ccl.highlight_row(1, 'YELLOW')
