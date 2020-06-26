"""Package Creator for CCL

Date: 2020-03-10
Revision: A
Author: Steven Fu
"""


import pandas as pd
from docx.api import Document
import re

from zipfile import ZipFile
import json
import copy

import shutil
import os


class Parser:
    """Parse the CCL into a DataFrame

    Notes:
        CCL will be used for the rest of the tool. Any edits to the filtered CCL will
        be reflected throughout the rest of the program.

    IMPORTANT!!!!!
    CCL HAS TO BE IN THE FORMAT OF THE 2020 TRF FORMAT (TRF REV M)
    *****

    Attrs:
        word_doc: the path to the word document
        document: the raw word document that was read
        table: the table object within the word document

    Functions:
        to_dataframe: Convert word directly to a dataframe
        filter: Will filter out the important information from the table
    """
    def __init__(self, word_doc):
        self.document = Document(word_doc)
        self.table = self.document.tables[0]

    def to_dataframe(self):
        """Convert word directly to dataframe with no regex/ filtering

        Notes:
            Bold heading are the bolded cable entries in the TRF. Will be mainly used to
            create the sub folders/ folder hierarchy for the CCL documents

        :return dataframe with columns: columns and data of the raw table document
        """
        columns = ['pn', 'desc', 'vendor', 'model', 'technical', 'standards', 'marks', 'bold']
        to_df = []
        for row in self.table.rows:
            temp = [cell.text.strip() for cell in row.cells]
            try:
                bold = True if row.cells[0].paragraphs[0].runs[0].bold else False
            except IndexError:  # Note when this happens it means that the text was a hyper link or not plain text
                print(f'Error occurred when parsing pn after {to_df[len(to_df)-1][0]}')
            to_df.append(temp + [bold])
        return pd.DataFrame(data=to_df, columns=columns)

    def filter(self):
        """Filter the inputted table into useful data only

        Note:
            Nums are ints where possible, the rest are string. All regex is done by the
            static method _re_get_cols. Please see for more information.

        :return DataFrame with columns: pn, desc, fn, assy, sch, bold.
        """
        def remove_duplicates(df):
            """Will only remove duplicates that are touching

            Main use is to remove/ clean up the original merged data.
            Input is the filtered dataframe.
            """
            pointer1, pointer2 = 0, 1
            while pointer2 < len(df.index):
                if df.loc[pointer1].equals(df.loc[pointer2]):
                    df = df.drop([pointer1])
                pointer1 = pointer2
                pointer2 += 1
            return df.reset_index(drop=True)

        columns = ['pn', 'desc', 'fn', 'dnums', 'illustration data']
        df = self.to_dataframe()
        data = [Parser._re_getcols(row) for index, row in df.iterrows()]
        filtered = pd.concat([pd.DataFrame(data, columns=columns), df['bold']], axis=1)
        return remove_duplicates(filtered)

    @staticmethod
    def _re_getcols(series: pd.Series):
        """Use regex to filter out data from series"""

        desc, fn = _re_fn_name(series['desc'])
        output = [
            _re_pn(series['pn']),
            desc,
            fn,
            _re_doc_num(series['technical']),
            illustration_dict(series['technical'])
        ]
        return output


def illustration_dict(string):
    """Creates a dictionary of all important data regarding illustrations

    returns [{'num': num, 'type': sch/assy, 'dnum': DXXXXXXX},...]
    """
    string = re.sub(r'\W+', '', string).replace(' ', '')
    if re.findall(r'refertoill.', string, re.IGNORECASE):
        ill_dict = [
            {'num': result[0], 'type': result[1], 'dnum': result[2]}
            for result in re.findall(r'(\d+)(assy|sch)(D\d+)', string, re.IGNORECASE)
        ]
        return ill_dict


def _re_pn(string):
    """Get the part number"""

    # Find any/ all digits with format of 6 concurrent numbers
    num = re.findall(r'\d+', string)
    return int(num[0]) if num else None


def _re_fn_name(string):
    """Find function number"""

    # Looks for (#X) as regex
    fn = re.findall(r'\(#(\d+)\)', string)
    if fn:
        to_pop = f'(#{fn[0]})'
        string = string.replace(to_pop, '').strip().replace('\n', '')
        # After removing the function number from string, needs to remove all
        # non A-Z 0-9 characters
        string = re.sub(r'([^\s\w]|_)+', '', string)
    return string, int(fn[0]) if fn else None


def _re_doc_num(string):
    """Finds document number, same logic as pn but with D in front

    used for verification purposes
    """
    dnums = re.findall(r'D\d+', string)
    return dnums if dnums else None


class Parent:
    """Unencodes the level information on the AVL multi-level BOM

    Note:
        The AVL multi level BOM must have atleast the "Level" and "Name" for this class to work

    Attributes:
        bom: the avl multi level bom as a DataFrame
        lowest: the lowest child (highest level) in the bom
        flat: the flattened dictionary
        tree: a multi level dictionary
    """
    def __init__(self, avl_bom):
        self.avl_bom = avl_bom
        self.set_bom()
        self.lowest = self.bom['Level'].max()
        self.flat = {}
        self.tree = {}

    def set_bom(self):
        """Sets self.bom depending if given a dataframe or path"""

        try:
            self.bom = pd.read_csv(self.avl_bom)
        except ValueError:
            if isinstance(self.avl_bom, pd.DataFrame):
                self.bom = self.avl_bom
            else:
                raise ValueError('Invalid AVL Bom input, needs to be path or dataframe')

    def build_flat(self):
        """Builds a flat dictionary with parent + child as keys

        Note:
            build_flat does not modify the original bom.

        :return flat: will assign the flat attribute a dictionary containing ALL parents as keys
                    in AVL bom with child listed as list
        """
        # Uses 2 pointers to check if next part is parent of top
        parent, child, parent_idx = 1, 2, None
        # Will check 1 level at a time starting with 1 ending with the lowest level
        while parent < self.lowest:
            # Will scan through entire BOM for every new level
            for idx in self.bom.index:

                if self.bom.loc[idx, 'Level'] == parent:
                    parent_idx = f'{idx} {self.bom.loc[idx, "Name"]}'
                    self.flat[parent_idx] = {}

                elif self.bom.loc[idx, 'Level'] == child and parent_idx is not None:
                    self.flat[parent_idx][f'{idx} {self.bom.loc[idx, "Name"]}'] = {}

            parent += 1  # Search for next level down parent
            child += 1
        return self.flat

    def build_tree(self):
        """Returns a structured tree using the flattened one

        Using the flat dictionary, it will reassemble it using the idea that there will be repeat
        keys. The child will now become the parent of the individual parent. For example, if
        12:{1,2,3,4}, and 15:{10, 11, 12, 13, 14} --> 15:{10, 11, 12:{1,2,3,4}, 13, 14}.

        After the child gets set, it will pop a key from a copy of the flattened list. Repeat until copied
        list is empty.

        :return dictionary with same list level as Enovia BOM
        """
        if not self.flat:
            self.build_flat()
        # Creates a copy of the flat attribute for popping
        copy_flat = copy.deepcopy(self.flat)
        while copy_flat:
            key = list(copy_flat)[0]
            # If empty, ie the first run, will set first key as first parent
            # This is ok because the flattened tree is in the same order as the bom.
            if not self.tree:
                self.tree[key] = copy_flat.pop(key)
                continue
            data = copy_flat.pop(key)
            Parent.replace_item(self.tree, key, data)
            # Checks if an item was inserted, if not the data will be appended to end
            if not Parent.exists(self.tree, key):
                self.tree[key] = data
        return self.tree

    @staticmethod
    def replace_item(obj, key, replace_value):
        """Recursively replaces item/ inserts item"""

        for k, v in obj.items():
            if v:
                obj[k] = Parent.replace_item(v, key, replace_value)
        if key in obj:
            obj[key] = replace_value
        return obj

    @staticmethod
    def exists(obj, key):
        """Recursively checks if item exists in dictionary"""

        if key in obj:
            return True
        for k, v in obj.items():
            if Parent.exists(v, key):
                return True


def BuildPackage(save_path: os.path, word_doc: os.path, avl_bom: os.path):
    """Builds the CCL package using Parser and Parent classes

    Parameters:
        :param save_path: the path/ location of where to save the CCL package
        :param word_doc: path/ location of where the CCL is saved (docx)
        :param avl_bom: path/ location of where the AVL multi level bom is saved (csv)

    :return a zip file at the location of the save path
    """
    if os.path.exists(save_path):
        raise FileExistsError('File with same name already exists in folder')

    parent = Parent(avl_bom)
    parse = Parser(word_doc)

    try:
        # Create a temporary folder where all files are waiting to be zipped
        # Temporary folder is created in the cwd
        os.makedirs('temp')
        parse.filter().to_csv(os.path.join('temp', 'filter.csv'))
        parse.document.save(os.path.join('temp', 'ccl.docx'))
        parent.bom.to_csv(os.path.join('temp', 'bom.csv'))
        with open(os.path.join('temp', 'flat.json'), 'w') as write:
            json.dump(parent.build_flat(), write, indent=4)
        with open(os.path.join('temp', 'tree.json'), 'w') as write:
            json.dump(parent.build_tree(), write, indent=4)
        # Zip the folder with all its contents
        with ZipFile(save_path, 'w') as package:
            files = os.listdir('temp')
            for file in files:
                package.write(os.path.join('temp', file), file)
    except Exception as e:
        raise e
    finally:
        # Remove the temporary folder
        shutil.rmtree('temp')


if __name__ == '__main__':
    ccl = Parser('ccl.docx')
    print(ccl.filter())